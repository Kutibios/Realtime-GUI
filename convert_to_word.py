from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import re

def convert_md_to_docx():
    # Yeni bir Word belgesi oluştur
    doc = Document()
    
    # Başlık stilini ayarla
    title = doc.add_heading('Basit Programlama Dili Derleyicisi Projesi Raporu', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Markdown dosyasını oku
    with open('Rapor.md', 'r', encoding='utf-8') as file:
        content = file.read()
    
    # Bölümleri ayır
    sections = content.split('## ')
    
    for section in sections[1:]:  # İlk bölümü atla (başlık)
        # Bölüm başlığını ve içeriğini ayır
        lines = section.split('\n', 1)
        if len(lines) > 1:
            heading = lines[0]
            content = lines[1]
            
            # Bölüm başlığını ekle
            doc.add_heading(heading, level=1)
            
            # Alt bölümleri işle
            sub_sections = content.split('### ')
            for sub_section in sub_sections[1:]:
                sub_lines = sub_section.split('\n', 1)
                if len(sub_lines) > 1:
                    sub_heading = sub_lines[0]
                    sub_content = sub_lines[1]
                    
                    # Alt bölüm başlığını ekle
                    doc.add_heading(sub_heading, level=2)
                    
                    # Kod bloklarını işle
                    code_blocks = re.findall(r'```(?:python|c)?\n(.*?)```', sub_content, re.DOTALL)
                    for code_block in code_blocks:
                        # Kod bloğunu ekle
                        p = doc.add_paragraph()
                        p.add_run(code_block.strip()).font.name = 'Consolas'
                    
                    # Normal metni işle
                    text_parts = re.split(r'```(?:python|c)?\n.*?```', sub_content, flags=re.DOTALL)
                    for text_part in text_parts:
                        if text_part.strip():
                            # Madde işaretli listeleri işle
                            if text_part.strip().startswith('- '):
                                for item in text_part.split('\n'):
                                    if item.strip().startswith('- '):
                                        p = doc.add_paragraph(style='List Bullet')
                                        p.add_run(item[2:].strip())
                            else:
                                p = doc.add_paragraph()
                                p.add_run(text_part.strip())
    
    # Belgeyi kaydet
    doc.save('Rapor.docx')

if __name__ == '__main__':
    convert_md_to_docx() 